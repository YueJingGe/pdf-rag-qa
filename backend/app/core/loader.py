"""Document loaders for PDF, DOCX, TXT, and Markdown files."""

from __future__ import annotations

import re
from pathlib import Path
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}


def _html_table_to_text(html: str) -> str:
    """Convert an HTML table to readable plain text rows.

    Example output:
      | 患者数量 | 42 | 高剂量 42 | 低剂量 42 | 高剂量 42 |
      | PASI75应答n（%）| 39 (92.9%) | 39 (92.9%) | 40 (95.2%) | 40 (95.2%) |
    """
    from html.parser import HTMLParser

    class TableParser(HTMLParser):
        def __init__(self):
            super().__init__()
            self.rows: list[list[str]] = []
            self.current_row: list[str] = []
            self.current_cell = ""
            self.in_cell = False

        def handle_starttag(self, tag, attrs):
            if tag in ("td", "th"):
                self.in_cell = True
                self.current_cell = ""
            elif tag == "tr":
                self.current_row = []

        def handle_endtag(self, tag):
            if tag in ("td", "th"):
                self.in_cell = False
                self.current_row.append(self.current_cell.strip())
            elif tag == "tr":
                if self.current_row:
                    self.rows.append(self.current_row)

        def handle_data(self, data):
            if self.in_cell:
                self.current_cell += data

    parser = TableParser()
    parser.feed(html)

    if not parser.rows:
        return html

    lines = []
    for row in parser.rows:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _convert_html_tables(text: str) -> str:
    """Find all HTML tables in text and replace them with plain text representation."""
    pattern = r'<html><body><table>.*?</table></body></html>'
    matches = list(re.finditer(pattern, text, re.DOTALL))

    if not matches:
        return text

    result = text
    for match in reversed(matches):
        table_text = _html_table_to_text(match.group())
        result = result[:match.start()] + "\n" + table_text + "\n" + result[match.end():]

    return result


def _load_docx(file_path: Path) -> list[Document]:
    """Load a DOCX file using python-docx, extracting paragraphs and tables."""
    import docx

    doc = docx.Document(str(file_path))
    parts: list[str] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag == "p":
            # Paragraph
            text = element.text or ""
            if text.strip():
                parts.append(text.strip())
        elif tag == "tbl":
            # Table - convert to pipe-separated text
            table_rows: list[str] = []
            for tr in element.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr"):
                cells: list[str] = []
                for tc in tr.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc"):
                    cell_text = "".join(
                        node.text or ""
                        for node in tc.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                    )
                    cells.append(cell_text.strip())
                if cells:
                    table_rows.append("| " + " | ".join(cells) + " |")
            if table_rows:
                parts.append("\n".join(table_rows))

    content = "\n\n".join(parts)
    return [Document(
        page_content=content,
        metadata={"source_filename": file_path.name, "page": 0},
    )]


def _ocr_pdf(file_path: Path) -> list[Document]:
    """OCR a scanned PDF using tesseract + pdf2image."""
    from pdf2image import convert_from_path
    import pytesseract

    images = convert_from_path(str(file_path), dpi=300)
    documents = []
    for i, image in enumerate(images):
        text = pytesseract.image_to_string(image, lang="chi_sim+eng")
        if text.strip():
            documents.append(Document(
                page_content=text.strip(),
                metadata={"source_filename": file_path.name, "page": i},
            ))
    return documents


def load_document(file_path: str | Path) -> list[Document]:
    """Load a document and return a list of LangChain Document objects."""
    file_path = Path(file_path)
    extension = file_path.suffix.lower()

    if extension == ".pdf":
        loader = PyPDFLoader(str(file_path))
        documents = loader.load()
        # Check if PDF is scanned (all pages empty) — fallback to OCR
        has_text = any(doc.page_content.strip() for doc in documents)
        if not has_text:
            documents = _ocr_pdf(file_path)
        for doc in documents:
            doc.metadata["source_filename"] = file_path.name
            if "page" not in doc.metadata:
                doc.metadata["page"] = 0
            doc.page_content = _convert_html_tables(doc.page_content)
        return documents
    elif extension in (".docx", ".doc"):
        return _load_docx(file_path)
    elif extension == ".md":
        loader = TextLoader(str(file_path), encoding="utf-8")
    elif extension == ".txt":
        loader = TextLoader(str(file_path), encoding="utf-8")
    else:
        raise ValueError(f"Unsupported file type: {extension}. Supported: {SUPPORTED_EXTENSIONS}")

    documents = loader.load()

    for doc in documents:
        doc.metadata["source_filename"] = file_path.name
        if "page" not in doc.metadata:
            doc.metadata["page"] = 0
        # Convert HTML tables to readable text
        doc.page_content = _convert_html_tables(doc.page_content)

    return documents
